import streamlit as st
import requests
import json
from docx import Document
import os
from functions import generar_desde_mercado, generar_desde_problema, generar_desde_azar, generar_prop_valor_usuario, generar_propvalor, generar_modelo_negocio, generar_pitchdeck

# define variables
problema = ""
propuesta_valor = ""
modelo_negocio = ""
pitch_deck = ""

def write_answers_to_doc(problema, propuesta_valor, modelo_negocio, pitch_deck):
    doc = Document()

    # Add each answer to the document
    doc.add_heading('Problema', level=1)
    doc.add_paragraph(problema)

    doc.add_heading('Propuesta de Valor', level=1)
    doc.add_paragraph(propuesta_valor)

    doc.add_heading('Modelo de Negocio', level=1)
    doc.add_paragraph(modelo_negocio)

    doc.add_heading('Pitch Deck', level=1)
    doc.add_paragraph(pitch_deck)

    # Save the document to a temporary location
    doc.save('/tmp/answers.docx')

# Claude functions
def create_text(prompt):
    api_url = "https://api.anthropic.com/v1/complete"
    headers = {
        "Content-Type": "application/json",
        "X-API-Key": st.secrets["API_KEY"]  # Use the API key from Streamlit's secrets
    }

    # Prepare the prompt for Claude
    conversation = f"Human: {prompt}\n\nAssistant:"

    # Define the body of the request
    body = {
        "prompt": conversation,
        "model": "claude-2.0",
        "max_tokens_to_sample": 100000,
        "temperature": 0.6,
        "stop_sequences": ["\n\nHuman:"]
    }

    # Make a POST request to the Claude API
    try:
        response = requests.post(api_url, headers=headers, data=json.dumps(body))
        response.raise_for_status()
    except requests.exceptions.HTTPError as errh:
        st.error(f"HTTP Error: {errh}")
    except requests.exceptions.ConnectionError as errc:
        st.error(f"Error Connecting: {errc}")
    except requests.exceptions.Timeout as errt:
        st.error(f"Timeout Error: {errt}")
    except requests.exceptions.RequestException as err:
        st.error(f"Something went wrong: {err}")
    except Exception as e:
        st.error(f"Unexpected error: {e}")

     # Extract Claude's response from the JSON response
    result = response.json()

    # Return Claude's response as a string
    return result['completion'].strip()


def app():

    st.title("Emprendetor v0")

    st.markdown("Este es un aplicativo diseñado para uso académico en la clase Emprendimiento e Innovación de la Universidad de los Andes por el profesor Camilo Serna Zamora")

    with st.container():
        col1, col2 = st.columns(2)

        col1.image('logo_uniandes.png')  
        col2.image('bobo.jpg')  

    with st.container():
        
        # Initialize session state variables if not already done
        if "result" not in st.session_state:
            st.session_state.result = ""
        if "prompts" not in st.session_state:
            st.session_state.prompts = ""
        
        st.markdown("Vamos a emprender en los tiempos de los LLMs")

        st.markdown("Tu emprendimiento/proyecto parte de:")

        option = st.selectbox('Selecciona una opción', ["",
                                                        "El interés de atender un mercado", 
                                                        "El interés de solucionar un problema particular", 
                                                        "No tengo aún nada definido"])
        
        if option == "El interés de atender un mercado":
            mercado_proyecto = st.text_input("¿Cuál mercado?")
            if st.button('Generar Mercado', key='boton_generar_mercado'):
                # Call your function here
                with st.spinner('Escribiendo...'):
                    # Create the answer
                    st.session_state.problema = create_text(generar_desde_mercado(mercado_proyecto))

                    # Display the result
                    st.write(st.session_state.problema)
                    st.session_state.container_1 = True


        elif option == "El interés de solucionar un problema particular":
            problema_proyecto = st.text_input("¿Cuál problema?")
            if st.button('Generar Problema', key='boton_generar_problema'):
                # Call your function here
                with st.spinner('Escribiendo...'):
                    # Create the answer
                    st.session_state.problema = create_text(generar_desde_problema(problema_proyecto))

                    # Display the result
                    st.write(st.session_state.problema)
                    st.session_state.container_1 = True


        elif option == "No tengo aún nada definido":
            if st.button('Generar Azar', key='boton_generar_azar'):
                # Call your function here
                with st.spinner('Escribiendo...'):
                    # Create the answer
                    st.session_state.problema = create_text(generar_desde_azar())

                    # Display the result
                    st.write(st.session_state.problema)
                    st.session_state.container_1 = True


    if 'container_1' in st.session_state and st.session_state.container_1:
        with st.container():
            st.markdown("¿Tienes definida la propuesta de valor?")
            option_propuesta_valor = st.selectbox('Selecciona una opción', ["", "Sí", "No"])

            if option_propuesta_valor == "Sí":
                propuesta_valor_proyecto = st.text_input("¿Cuál es?")
                if st.button('Generar Propuesta de Valor', key='boton_generar_propuesta_valor'):
                    # Call your function here
                    with st.spinner('Escribiendo...'):
                        # Create the answer

                        st.session_state.propuesta_valor = create_text(generar_prop_valor_usuario(st.session_state.problema, propuesta_valor_proyecto))

                        # Display the result
                        st.write(st.session_state.propuesta_valor)
                        st.session_state.container_2 = True
                        

            elif option_propuesta_valor == "No":
                if st.button('Generar Propuesta de Valor sin Definir', key='boton_generar_propuesta_valor_0'):
                    # Call your function here
                    with st.spinner('Escribiendo...'):
                        # Create the answer
                        st.session_state.propuesta_valor = create_text(generar_propvalor(st.session_state.problema))

                        # Display the result
                        st.write(st.session_state.propuesta_valor)
                        st.session_state.container_2 = True

                   

        if 'container_2' in st.session_state and st.session_state.container_2:
            with st.container():
                st.markdown("Generar un modelo de negocio")
                if st.button('Generar Modelo de Negocio', key='boton_generar_modelo_negocio'):
                    with st.spinner('Escribiendo...'):
                        # Create the answer
                        st.session_state.modelo_negocio = create_text(generar_modelo_negocio(st.session_state.problema, st.session_state.propuesta_valor))

                        # Display the result
                        st.write(st.session_state.modelo_negocio)
                        st.session_state.container_3 = True


            if 'container_3' in st.session_state and st.session_state.container_3:
                with st.container():
                    st.markdown("Generar un pitch deck")
                    if st.button('Generar Pitch Deck', key='boton_generar_pitch_deck'):
                        with st.spinner('Escribiendo...'):
                            # Create the answer
                            st.session_state.pitch_deck = create_text(generar_pitchdeck(st.session_state.problema, st.session_state.propuesta_valor, st.session_state.modelo_negocio))

                            # Display the result
                            st.write(st.session_state.pitch_deck)
                        
if st.button('Descargar negocio'):
        with st.spinner('Preparando...'):
            # Call the function to write the answers to a .doc file
            write_answers_to_doc(st.session_state.problema, st.session_state.propuesta_valor, st.session_state.modelo_negocio, st.session_state.pitch_deck)

            # Provide a download link
            st.markdown('<a href="/tmp/answers.docx" download="answers.docx">Click here to download the .doc file</a>', unsafe_allow_html=True)         

        
if __name__ == "__main__":
    app()
